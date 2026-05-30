# -*- coding: utf-8 -*-
"""Build a single RTL Hebrew DOCX study book from study_content.json.
Pure learning content (sections only) — no quizzes/exercises."""
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
DATA = os.path.join(HERE, "study_content.json")
OUT = os.path.join(HERE, "..", "חומרי_לימוד", "מערכת_פוליטית_ישראלית__חומר_לימוד_מלא.docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
HEB_FONT = "David"
LAT_FONT = "Times New Roman"

def _rtl_par(p):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)

def _style_run(run, size, bold=False, italic=False, color=None):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic
    run.font.name = LAT_FONT
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rfonts = rPr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rPr.append(rfonts)
    rfonts.set(qn('w:cs'), HEB_FONT); rfonts.set(qn('w:ascii'), LAT_FONT); rfonts.set(qn('w:hAnsi'), LAT_FONT)
    rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)

def para(doc, text, size=12, bold=False, italic=False, color=None,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6, space_before=0,
         line=1.5, left_indent=None, right_indent=None):
    p = doc.add_paragraph()
    _rtl_par(p)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if left_indent is not None: pf.left_indent = Inches(left_indent)
    if right_indent is not None: pf.right_indent = Inches(right_indent)
    if text:
        _style_run(p.add_run(text), size, bold=bold, italic=italic, color=color)
    return p

def hrule(doc):
    p = doc.add_paragraph(); _rtl_par(p)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '1F3A5F')
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(4)

def comparison_table(doc, c):
    rows = c["rows"]
    if not rows: return
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tbl.style = "Table Grid"
    # force RTL visual order
    tblPr = tbl._tbl.tblPr
    bidi = OxmlElement('w:bidiVisual'); tblPr.append(bidi)
    hdr = tbl.rows[0].cells
    def cell(cellobj, title, sub, bold=True):
        cellobj.text = ""
        p = cellobj.paragraphs[0]; _rtl_par(p); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _style_run(p.add_run(title), 11, bold=bold, color=NAVY)
        if sub:
            p2 = cellobj.add_paragraph(); _rtl_par(p2); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _style_run(p2.add_run(sub), 9, italic=True, color=GREY)
    cell(hdr[0], "היבט", "")
    cell(hdr[1], c["rightTitle"], c["rightSubtitle"])
    cell(hdr[2], c["leftTitle"], c["leftSubtitle"])
    for r in rows:
        cells = tbl.add_row().cells
        for cobj, txt, b in ((cells[0], r["axis"], True), (cells[1], r["right"], False), (cells[2], r["left"], False)):
            cobj.text = ""
            p = cobj.paragraphs[0]; _rtl_par(p); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _style_run(p.add_run(txt or ""), 10.5, bold=b)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def main():
    data = json.load(open(DATA, encoding="utf-8"))
    doc = Document()
    # base style
    st = doc.styles["Normal"]; st.font.name = LAT_FONT; st.font.size = Pt(12)
    for s in doc.sections:
        s.right_margin = Inches(0.9); s.left_margin = Inches(0.9)

    # cover
    para(doc, "מערכת פוליטית ישראלית", size=26, bold=True, color=NAVY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=40, space_after=4)
    para(doc, "חומר לימוד מלא — תוכן כל היחידות", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, "פרופ' דורון נבות · אוניברסיטת חיפה", size=12, italic=True, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "מבוסס על תוכן הסקציות שבאתר political-system-app — ללא שאלות ותרגול.",
         size=10, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # table of contents
    para(doc, "תוכן עניינים", size=15, bold=True, color=NAVY, space_before=6, space_after=8)
    for u in data:
        para(doc, f"יחידה {u['number']} — {u['title']}", size=12, bold=True, space_after=2, right_indent=0.2)
    doc.add_page_break()

    for ui, u in enumerate(data):
        # unit header
        para(doc, f"יחידה {u['number']}", size=13, bold=True, color=GREY, space_after=0, space_before=4)
        para(doc, u["title"], size=20, bold=True, color=NAVY, space_after=4)
        if u.get("subtitle"):
            para(doc, u["subtitle"], size=11, italic=True, color=GREY, space_after=8)
        hrule(doc)
        if u.get("articles"):
            para(doc, "מאמרי חובה", size=12, bold=True, color=NAVY, space_before=4, space_after=2)
            for a in u["articles"]:
                para(doc, "• " + a, size=11, space_after=1, right_indent=0.2)
        if u.get("leadQuestion"):
            para(doc, "שאלת המנחה", size=12, bold=True, color=NAVY, space_before=6, space_after=2)
            para(doc, u["leadQuestion"], size=11, italic=True, space_after=8, right_indent=0.2)
        if u.get("objectives"):
            para(doc, "מטרות הלמידה", size=12, bold=True, color=NAVY, space_before=4, space_after=2)
            for o in u["objectives"]:
                para(doc, "• " + o, size=11, space_after=2, right_indent=0.2)

        for g in u["groups"]:
            if g.get("partTitle"):
                para(doc, g["partTitle"], size=15, bold=True, color=NAVY, space_before=14, space_after=2)
                if g.get("partSubtitle"):
                    para(doc, g["partSubtitle"], size=10.5, italic=True, color=GREY, space_after=6)
            for sec in g["sections"]:
                para(doc, sec["heading"], size=13.5, bold=True, color=NAVY,
                     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=4)
                if sec.get("tldr"):
                    para(doc, "בקצרה: " + sec["tldr"], size=11, italic=True, color=GREY,
                         space_after=6, right_indent=0.15, left_indent=0.15)
                for p in sec["paragraphs"]:
                    para(doc, p, size=12, space_after=6)
                if sec.get("comparison"):
                    c = sec["comparison"]
                    if c.get("title"):
                        para(doc, c["title"], size=11.5, bold=True, color=NAVY, space_before=4, space_after=3,
                             align=WD_ALIGN_PARAGRAPH.CENTER)
                    comparison_table(doc, c)
                if sec.get("quote"):
                    q = sec["quote"]
                    para(doc, "“" + q["text"] + "”", size=11.5, italic=True, color=NAVY,
                         space_before=4, space_after=2, right_indent=0.4, left_indent=0.4)
                    src = " — ".join([x for x in [q.get("source"), (("עמ' " + str(q["page"])) if q.get("page") else None)] if x])
                    if src:
                        para(doc, src, size=10, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.LEFT,
                             space_after=6, right_indent=0.4, left_indent=0.4)
                if sec.get("callout"):
                    co = sec["callout"]
                    title = (co.get("title") or "").strip()
                    body = co["text"]
                    p = para(doc, "", size=11, space_before=4, space_after=8, right_indent=0.15, left_indent=0.15)
                    if title:
                        _style_run(p.add_run(title + "  "), 11, bold=True, color=NAVY)
                    _style_run(p.add_run(body), 11, italic=False)
                    # light shading box
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),'F2F5FA')
                    pPr.append(shd)
        if ui != len(data) - 1:
            doc.add_page_break()

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("Saved:", os.path.abspath(OUT))

if __name__ == "__main__":
    main()
